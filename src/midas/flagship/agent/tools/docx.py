"""docx.draft — Word document artifact, gated.

Behind the ``[docs]`` extra (``python-docx``). Without the extra, the planner
returns a clear error rather than silently degrading; the artifact factory's
``artifact.text`` fallback covers the never-refuse contract.
"""

from __future__ import annotations

import hashlib
from dataclasses import dataclass

from .fsguard import FsGuard


@dataclass(frozen=True)
class DocxPlan:
    path: str
    bytes_len: int
    sha256_new: str
    sha256_prev: str | None
    preview: str


def _sha256(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def plan_docx(guard: FsGuard, path: str, *, title: str, body: str) -> DocxPlan:
    target = guard.resolve(path)
    if target.suffix.lower() != ".docx":
        raise ValueError(f"docx path must end in .docx, got {target.suffix!r}")
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - exercised via the [docs] extra
        raise RuntimeError(
            "docx.draft requires the [docs] extra (python-docx)"
        ) from exc

    import io

    document = Document()
    document.add_heading(title, level=1)
    for paragraph in body.split("\n\n"):
        document.add_paragraph(paragraph)
    buf = io.BytesIO()
    document.save(buf)
    data = buf.getvalue()
    sha_prev = None
    if target.exists() and target.is_file():
        sha_prev = _sha256(target.read_bytes())
    return DocxPlan(
        path=str(target),
        bytes_len=len(data),
        sha256_new=_sha256(data),
        sha256_prev=sha_prev,
        preview=f"{title}\n\n{body[:400]}",
    )


def execute_docx(guard: FsGuard, *, path: str, title: str, body: str) -> DocxPlan:
    """Materialize the docx after approval."""
    target = guard.resolve(path)
    plan = plan_docx(guard, path, title=title, body=body)
    target.parent.mkdir(parents=True, exist_ok=True)
    # Re-render so the on-disk bytes match plan.sha256_new exactly.
    from docx import Document

    document = Document()
    document.add_heading(title, level=1)
    for paragraph in body.split("\n\n"):
        document.add_paragraph(paragraph)
    document.save(target)
    return plan
